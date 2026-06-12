from flask import Flask, render_template, request, send_file, send_from_directory, jsonify
import os
import io
import uuid
import subprocess
import tempfile
from werkzeug.utils import secure_filename
import secrets
import json
import re

# Progress tracking for real-time conversion feedback
import threading
_progress = {}
_progress_lock = threading.Lock()

def _set_progress(task_id, percent, message=""):
    with _progress_lock:
        _progress[task_id] = {"percent": percent, "message": message}

def _get_progress(task_id):
    with _progress_lock:
        return _progress.get(task_id)

def _clear_progress(task_id):
    with _progress_lock:
        _progress.pop(task_id, None)


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500MB for video
app.config["UPLOAD_FOLDER"] = "/tmp/fileforge_uploads"

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(os.path.join(app.config["UPLOAD_FOLDER"], "tmp"), exist_ok=True)
DOWNLOAD_DIR = "/tmp/fileforge_downloads"
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

FFMPEG = "/usr/bin/ffmpeg"
FFPROBE = "/usr/bin/ffprobe"

# Preloaded font cache for maximum speed
from fpdf import FPDF
_FONT_PATH = '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
_FONT_OK = os.path.exists(_FONT_PATH)
if _FONT_OK:
    _FONT_OK = FPDF()
    _FONT_OK.add_page()
    _FONT_OK.add_font("F", "", _FONT_PATH)
    _FONT_OK.set_font("F", "", 11)
else:
    _FONT_OK = None


# ── Conversion type definitions (used by both frontend and backend) ──

CONVERSIONS = {
    # Documents
    "pdf2docx":   {"label": "PDF → Word", "icon": "📝", "cat": "doc", "in": ".pdf", "hint": "PDF 转可编辑 Word"},
    "docx2pdf":   {"label": "Word → PDF", "icon": "📋", "cat": "doc", "in": ".docx", "hint": "Word 文档转 PDF"},
    "pdf2txt":    {"label": "PDF → 文本", "icon": "📄", "cat": "doc", "in": ".pdf", "hint": "提取 PDF 中的文字"},
    "docx2txt":   {"label": "Word → 文本", "icon": "📃", "cat": "doc", "in": ".docx", "hint": "提取 Word 中的文字"},
    "md2pdf":     {"label": "Markdown→PDF", "icon": "📘", "cat": "doc", "in": ".md", "hint": "Markdown 转 PDF"},
    "html2pdf":   {"label": "HTML → PDF", "icon": "🌐", "cat": "doc", "in": ".html", "hint": "网页/HTML 转 PDF"},
    "txt2pdf":    {"label": "文本 → PDF", "icon": "📕", "cat": "doc", "in": ".txt", "hint": "纯文本转 PDF"},

    # Images
    "img2png":    {"label": "图片 → PNG", "icon": "🖼️", "cat": "img", "in": ".jpg,.jpeg,.png,.webp,.bmp,.gif,.tiff,.ico", "hint": "任意图片转 PNG"},
    "img2jpg":    {"label": "图片 → JPG", "icon": "📸", "cat": "img", "in": ".jpg,.jpeg,.png,.webp,.bmp,.gif,.tiff,.ico", "hint": "任意图片转 JPG"},
    "img2webp":   {"label": "图片 → WebP", "icon": "🌐", "cat": "img", "in": ".jpg,.jpeg,.png,.webp,.bmp,.gif,.tiff,.ico", "hint": "任意图片转 WebP"},
    "img2bmp":    {"label": "图片 → BMP", "icon": "🟦", "cat": "img", "in": ".jpg,.jpeg,.png,.webp,.bmp,.gif,.tiff,.ico", "hint": "任意图片转 BMP"},
    "img2ico":    {"label": "图片 → ICO", "icon": "🔷", "cat": "img", "in": ".jpg,.jpeg,.png,.webp,.bmp,.gif,.tiff", "hint": "生成图标文件"},
    "img2tiff":   {"label": "图片 → TIFF", "icon": "🟩", "cat": "img", "in": ".jpg,.jpeg,.png,.webp,.bmp,.gif,.ico", "hint": "无损 TIFF 格式"},

    # Audio
    "audio2mp3":  {"label": "音频 → MP3", "icon": "🎵", "cat": "audio", "in": ".mp3,.wav,.ogg,.flac,.m4a,.aac,.wma,.opus,.aiff", "hint": "任意音频转 MP3"},
    "audio2wav":  {"label": "音频 → WAV", "icon": "🎶", "cat": "audio", "in": ".mp3,.wav,.ogg,.flac,.m4a,.aac,.wma,.opus,.aiff", "hint": "任意音频转 WAV 无损"},
    "audio2ogg":  {"label": "音频 → OGG", "icon": "🎼", "cat": "audio", "in": ".mp3,.wav,.ogg,.flac,.m4a,.aac,.wma,.opus,.aiff", "hint": "任意音频转 OGG"},
    "audio2flac": {"label": "音频 → FLAC", "icon": "🎹", "cat": "audio", "in": ".mp3,.wav,.ogg,.flac,.m4a,.aac,.wma,.opus,.aiff", "hint": "任意音频转 FLAC 无损"},
    "audio2aac":  {"label": "音频 → AAC", "icon": "🎧", "cat": "audio", "in": ".mp3,.wav,.ogg,.flac,.m4a,.aac,.wma,.opus,.aiff", "hint": "任意音频转 AAC"},

    # Video to Audio
    "video2mp3":  {"label": "视频 → MP3", "icon": "🎬→🎵", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "提取视频中的音频"},
    "video2wav":  {"label": "视频 → WAV", "icon": "🎬→🎶", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "提取视频音频为 WAV"},

    # Video to GIF
    "video2gif":  {"label": "视频 → GIF", "icon": "🎬→🎞️", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "视频转 GIF 动图"},

    # Video Format
    "video2mp4":  {"label": "视频 → MP4", "icon": "🎬", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "任意视频转 MP4 (H.264)"},
    "video2avi":  {"label": "视频 → AVI", "icon": "📀", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "任意视频转 AVI"},
    "video2mov":  {"label": "视频 → MOV", "icon": "🍎", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "任意视频转 MOV"},
    "video2webm": {"label": "视频 → WebM", "icon": "🌍", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "任意视频转 WebM"},
    "video2mkv":  {"label": "视频 → MKV", "icon": "🎦", "cat": "video", "in": ".mp4,.avi,.mov,.mkv,.webm,.flv,.wmv,.m4v", "hint": "任意视频转 MKV"},

    # Spreadsheet
    "csv2xlsx":   {"label": "CSV → Excel", "icon": "📊", "cat": "sheet", "in": ".csv", "hint": "CSV 转 Excel 表格"},
    "xlsx2csv":   {"label": "Excel → CSV", "icon": "📈", "cat": "sheet", "in": ".xlsx,.xls", "hint": "Excel 转 CSV"},
}

# ── Helpers ──

def _save_temp(file, suffix):
    suffix = secure_filename(suffix) or ".tmp"
    fname = str(uuid.uuid4()) + suffix
    path = os.path.join("/tmp", "fileforge_" + fname)
    file.save(path)
    return path
def _safe_unlink(path):
    try:
        if os.path.exists(path):
            os.unlink(path)
    except Exception:
        pass

def _run_ffmpeg(args, task_id=None, duration=None):
    """Run ffmpeg with args and real-time progress. Raise on failure."""
    cmd = [FFMPEG, "-y", "-hide_banner", "-loglevel", "error", "-progress", "pipe:2", "-nostats"] + args
    proc = subprocess.Popen(cmd, stderr=subprocess.PIPE, text=True, bufsize=1)
    last_pct = 0
    try:
        for line in proc.stderr:
            line = line.strip()
            if line.startswith("out_time_ms=") and task_id and duration and duration > 0:
                try:
                    ms = int(line.split("=")[1])
                    pct = min(95, max(5, int(ms / (duration * 1000) * 100)))
                    if pct > last_pct + 2:
                        last_pct = pct
                        _set_progress(task_id, pct, "转码中...")
                except Exception:
                    pass
            elif line.startswith("out_time=") and task_id and duration and duration > 0:
                try:
                    parts = line.split("=")[1].split(":")
                    secs = int(parts[0])*3600 + int(parts[1])*60 + float(parts[2])
                    pct = min(95, max(5, int(secs / duration * 100)))
                    if pct > last_pct + 2:
                        last_pct = pct
                        _set_progress(task_id, pct, "转码中...")
                except Exception:
                    pass
    finally:
        proc.wait()
    if proc.returncode != 0:
        remaining = proc.stderr.read() if not proc.stderr.closed else ""
        raise RuntimeError(remaining.strip() or f"ffmpeg exit {proc.returncode}")
    return proc


def _get_duration(path):
    """Get media duration in seconds using ffprobe."""
    try:
        result = subprocess.run(
            [FFMPEG.replace("ffmpeg", "ffprobe"), "-v", "error", "-show_entries",
             "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", path],
            capture_output=True, text=True, timeout=30
        )
        return float(result.stdout.strip())
    except Exception:
        return None

def _download_name(original, new_ext):
    base = secure_filename(original.rsplit(".", 1)[0] if "." in original else original)
    if not base:
        base = "converted"
    return base + new_ext

# ── Simple rate limiter ──
from functools import wraps
from time import time
_rate_limits = {}
def rate_limit(max_per_minute=20):
    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            ip = request.headers.get("X-Real-IP", request.remote_addr) or "unknown"
            now = time()
            records = _rate_limits.get(ip, [])
            records = [t for t in records if now - t < 60]
            if len(records) >= max_per_minute:
                return jsonify({"error": "请求过于频繁，请稍后再试"}), 429
            records.append(now)
            _rate_limits[ip] = records
            return f(*args, **kwargs)
        return wrapped
    return decorator

# ── Routes ──

@app.route("/")
def index():
    return render_template("index.html", conversions=CONVERSIONS)


@app.route("/batch")
def batch_page():
    return render_template("batch.html")

@app.route("/api/conversions")
def api_conversions():
    return jsonify(CONVERSIONS)

@app.route("/convert/progress/<task_id>")
def convert_progress(task_id):
    p = _get_progress(task_id)
    if p is None:
        return jsonify({"percent": -1})
    return jsonify(p)

@app.route("/convert", methods=["POST"])
@rate_limit(max_per_minute=20)
def convert():
    if "file" not in request.files:
        return jsonify({"error": "no file"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "empty"}), 400
    conv_type = request.form.get("type", "")
    if conv_type not in CONVERSIONS:
        return jsonify({"error": "unsupported"}), 400
    task_id = request.form.get("task_id", str(uuid.uuid4()))
    _set_progress(task_id, 0, "preparing")
    try:
        result = _do_convert(file, conv_type, task_id)
        download_token = secrets.token_urlsafe(32)
        result.direct_passthrough = False
        data = result.get_data()
        dest = os.path.join(DOWNLOAD_DIR, download_token)
        with open(dest, "wb") as f:
            f.write(data)
        dname = "converted_file"
        disposition = result.headers.get("Content-Disposition", "")
        match = re.search(r"filename[^=]*=([^;]+)", disposition)
        if match:
            dname = match.group(1)
        mime = result.mimetype or "application/octet-stream"
        with open(dest + ".json", "w") as f:
            json.dump({"name": dname, "mime": mime}, f)
        _set_progress(task_id, 100, "done")
        return jsonify({"ok": True, "download": f"/download/{download_token}", "name": dname})
    except subprocess.TimeoutExpired:
        _clear_progress(task_id)
        return jsonify({"error": "timeout"}), 500
    except Exception as e:
        _clear_progress(task_id)
        return jsonify({"error": "conversion failed"}), 500

def _do_convert(file, conv_type, task_id=None):
    # ── Document conversions ──
    if conv_type == "pdf2docx":
        return _pdf2docx(file, task_id)
    elif conv_type == "docx2pdf":
        return _docx2pdf(file, task_id)
    elif conv_type == "pdf2txt":
        return _pdf2txt(file, task_id)
    elif conv_type == "docx2txt":
        return _docx2txt(file, task_id)
    elif conv_type == "md2pdf":
        return _md2pdf(file, task_id)
    elif conv_type == "html2pdf":
        return _html2pdf(file, task_id)
    elif conv_type == "txt2pdf":
        return _txt2pdf(file, task_id)

    # ── Image conversions ──
    elif conv_type.startswith("img2"):
        fmt_map = {"img2png": "PNG", "img2jpg": "JPEG", "img2webp": "WEBP", "img2bmp": "BMP", "img2ico": "ICO", "img2tiff": "TIFF"}
        return _convert_image(file, fmt_map[conv_type])

    # ── Audio conversions ──
    elif conv_type.startswith("audio2"):
        fmt_map = {"audio2mp3": "mp3", "audio2wav": "wav", "audio2ogg": "ogg", "audio2flac": "flac", "audio2aac": "aac"}
        return _convert_audio(file, task_id, fmt_map[conv_type])

    # ── Video to Audio ──
    elif conv_type in ("video2mp3", "video2wav"):
        fmt = "mp3" if conv_type == "video2mp3" else "wav"
        return _video_extract_audio(file, task_id, fmt)

    # ── Video to GIF ──
    elif conv_type == "video2gif":
        return _video2gif(file, task_id)

    # ── Video Format ──
    elif conv_type.startswith("video2"):
        fmt_map = {"video2mp4": "mp4", "video2avi": "avi", "video2mov": "mov", "video2webm": "webm", "video2mkv": "mkv"}
        return _convert_video(file, task_id, fmt_map[conv_type])

    # ── Spreadsheet ──
    elif conv_type == "csv2xlsx":
        return _csv2xlsx(file, task_id)
    elif conv_type == "xlsx2csv":
        return _xlsx2csv(file, task_id)

    return jsonify({"error": "未知转换类型"}), 400

# ── Document Converters ──

def _pdf2docx(file, task_id):
    if task_id: _set_progress(task_id, 10, '解析PDF...')
    from pdf2docx import Converter
    pdf_path = _save_temp(file, ".pdf")
    docx_path = pdf_path.rsplit(".", 1)[0] + ".docx"
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        return send_file(docx_path, as_attachment=True,
                        download_name=_download_name(file.filename, ".docx"),
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    finally:
        _safe_unlink(pdf_path)
        _safe_unlink(docx_path)

def _docx2pdf(file):
    from docx import Document
    from fpdf import FPDF
    docx_path = _save_temp(file, ".docx")
    try:
        doc = Document(docx_path)
        pdf = FPDF()
        pdf.add_page()
        font_path = _FONT_PATH
        if os.path.exists(font_path):
            pdf.add_font("F", "", font_path, uni=True)
            pdf.set_font("F", "", 11)
        else:
            pdf.set_font("Helvetica", "", 11)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                pdf.multi_cell(0, 7, text[:500])
                pdf.ln(2)
        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True,
                        download_name=_download_name(file.filename, ".pdf"),
                        mimetype="application/pdf")
    finally:
        _safe_unlink(docx_path)

def _pdf2txt(file, task_id):
    if task_id: _set_progress(task_id, 20, '提取文字...')
    import pdfplumber
    with pdfplumber.open(file.stream) as pdf:
        text = "\n\n".join(page.extract_text() or "" for page in pdf.pages)
    buf = io.BytesIO(text.encode("utf-8"))
    return send_file(buf, as_attachment=True,
                    download_name=_download_name(file.filename, ".txt"),
                    mimetype="text/plain")

def _docx2txt(file, task_id):
    if task_id: _set_progress(task_id, 20, '提取文字...')
    from docx import Document
    doc = Document(file.stream)
    text = "\n".join(p.text for p in doc.paragraphs)
    buf = io.BytesIO(text.encode("utf-8"))
    return send_file(buf, as_attachment=True,
                    download_name=_download_name(file.filename, ".txt"),
                    mimetype="text/plain")




def _md2pdf(file, task_id):
    import markdown as mdlib
    import re; from fpdf import FPDF
    md_text = file.stream.read().decode("utf-8")
    text = re.sub(r"<[^>]+>", "", mdlib.markdown(md_text))
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("F", "", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", uni=True)
    pdf.set_font("F", "", 11)
    for line in text.split("\n"):
        pdf.cell(0, 7, line[:200], new_x="LMARGIN", new_y="NEXT")
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                    download_name=_download_name(file.filename, ".pdf"),
                    mimetype="application/pdf")

def _html2pdf(file, task_id):
    import re; from fpdf import FPDF
    text = re.sub(r"<[^>]+>", "", file.stream.read().decode("utf-8"))
    text = re.sub(r"\n{3,}", "\n\n", text)
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("F", "", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", uni=True)
    pdf.set_font("F", "", 11)
    for line in text.split("\n"):
        if line.strip():
            pdf.cell(0, 7, line[:200], new_x="LMARGIN", new_y="NEXT")
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                    download_name=_download_name(file.filename, ".pdf"),
                    mimetype="application/pdf")

def _txt2pdf(file, task_id=None):
    text = file.stream.read().decode('utf-8')
    if task_id: _set_progress(task_id, 10, '读取文本...')
    if _FONT_OK:
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font("F", "", _FONT_PATH)
        pdf.set_font("F", "", 11)
    else:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "", 11)
    if task_id: _set_progress(task_id, 30, '生成PDF...')
    pdf.multi_cell(0, 6, text[:50000])
    if task_id: _set_progress(task_id, 70, '输出文件...')
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                    download_name=_download_name(file.filename, '.pdf'),
                    mimetype='application/pdf')

def _convert_image(file, target_format):
    from PIL import Image
    img = Image.open(file.stream)
    buf = io.BytesIO()
    save_kwargs = {}
    if target_format == "JPEG":
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        save_kwargs["quality"] = 92
    elif target_format == "WEBP":
        save_kwargs["quality"] = 85
    elif target_format == "ICO":
        img = img.resize((max(16, min(256, img.width)), max(16, min(256, img.height))))
        if img.mode != "RGBA":
            img = img.convert("RGBA")
    elif target_format == "TIFF":
        save_kwargs["compression"] = "tiff_lzw"
    ext_map = {"JPEG": "jpg", "PNG": "png", "WEBP": "webp", "BMP": "bmp", "ICO": "ico", "TIFF": "tiff"}
    mime_map = {"JPEG": "image/jpeg", "PNG": "image/png", "WEBP": "image/webp", "BMP": "image/bmp", "ICO": "image/x-icon", "TIFF": "image/tiff"}
    img.save(buf, format=target_format, **save_kwargs)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                    download_name=_download_name(file.filename, "." + ext_map[target_format]),
                    mimetype=mime_map[target_format])

# ── Audio Converter (ffmpeg) ──

FFMPEG_AUDIO_CODECS = {
    "mp3":  ["-threads", "2", "-codec:a", "libmp3lame", "-b:a", "192k", "-q:a", "2"],
    "wav":  ["-threads", "2", "-codec:a", "pcm_s16le"],
    "ogg":  ["-threads", "2", "-codec:a", "libvorbis", "-b:a", "192k", "-q:a", "3"],
    "flac": ["-threads", "2", "-codec:a", "flac", "-compression_level", "8"],
    "aac":  ["-threads", "2", "-codec:a", "aac", "-b:a", "192k"],
}

def _convert_audio(file, task_id, fmt):
    in_path = _save_temp(file, os.path.splitext(file.filename)[1] or ".tmp")
    out_path = in_path.rsplit(".", 1)[0] + "_out." + fmt
    in_ext = os.path.splitext(file.filename)[1].lower()
    codec_args = FFMPEG_AUDIO_CODECS[fmt]
    try:
        if task_id: _set_progress(task_id, 3, '分析音频...')
        # Same format? Stream copy
        if in_ext == '.' + fmt:
            if task_id: _set_progress(task_id, 10, '快速封装(无需转码)...')
            _run_ffmpeg(["-i", in_path, "-c", "copy", out_path])
        else:
            dur = _get_duration(in_path)
            if task_id: _set_progress(task_id, 5, '开始转码...')
            _run_ffmpeg(["-i", in_path] + codec_args + [out_path], task_id=task_id, duration=dur)
        mimes = {"mp3": "audio/mpeg", "wav": "audio/wav", "ogg": "audio/ogg", "flac": "audio/flac", "aac": "audio/aac"}
        return send_file(out_path, as_attachment=True,
                        download_name=_download_name(file.filename, "." + fmt),
                        mimetype=mimes[fmt])
    finally:
        _safe_unlink(in_path)
        _safe_unlink(out_path)

# ── Video Extract Audio ──

def _video_extract_audio(file, task_id, fmt):
    in_path = _save_temp(file, os.path.splitext(file.filename)[1] or ".mp4")
    out_path = in_path.rsplit(".", 1)[0] + "_out." + fmt
    try:
        if task_id: _set_progress(task_id, 3, '分析媒体...')
        dur = _get_duration(in_path)
        if task_id: _set_progress(task_id, 5, '提取音频...')
        if fmt == "mp3":
            _run_ffmpeg(["-i", in_path, "-vn", "-codec:a", "libmp3lame", "-b:a", "192k", out_path], task_id=task_id, duration=dur)
        else:
            _run_ffmpeg(["-i", in_path, "-vn", "-codec:a", "pcm_s16le", out_path], task_id=task_id, duration=dur)
        mimes = {"mp3": "audio/mpeg", "wav": "audio/wav"}
        return send_file(out_path, as_attachment=True,
                        download_name=_download_name(file.filename, "." + fmt),
                        mimetype=mimes[fmt])
    finally:
        _safe_unlink(in_path)
        _safe_unlink(out_path)

# ── Video to GIF ──

def _video2gif(file, task_id):
    in_path = _save_temp(file, os.path.splitext(file.filename)[1] or ".mp4")
    palette_path = in_path.rsplit(".", 1)[0] + "_palette.png"
    out_path = in_path.rsplit(".", 1)[0] + ".gif"
    try:
        if task_id: _set_progress(task_id, 3, '分析视频...')
        dur = _get_duration(in_path)
        if task_id: _set_progress(task_id, 5, '生成调色板...')
        _run_ffmpeg(["-i", in_path, "-vf", "fps=10,scale=480:-1:flags=lanczos,palettegen", palette_path])
        if task_id: _set_progress(task_id, 50, '生成GIF...')
        _run_ffmpeg(["-i", in_path, "-i", palette_path, "-filter_complex", "fps=10,scale=480:-1:flags=lanczos[x];[x][1:v]paletteuse", out_path], task_id=task_id, duration=dur)
        return send_file(out_path, as_attachment=True,
                        download_name=_download_name(file.filename, ".gif"),
                        mimetype="image/gif")
    finally:
        _safe_unlink(in_path)
        _safe_unlink(palette_path)
        _safe_unlink(out_path)

# ── Video Format Converter (ffmpeg) ──

def _is_same_codec_family(in_ext, out_fmt):
    """Check if we can stream-copy instead of re-encode."""
    compatible = {
        ".mp4": ["mp4", "mov", "mkv"],
        ".mov": ["mp4", "mov", "mkv"],
        ".mkv": ["mp4", "mov", "mkv"],
        ".avi": ["avi"],
        ".webm": ["webm"],
    }
    return out_fmt in compatible.get(in_ext.lower(), [])

FFMPEG_VIDEO_FORMATS = {
    "mp4":  ["-threads", "2", "-c:v", "libx264", "-preset", "ultrafast", "-tune", "fastdecode", "-crf", "23", "-c:a", "aac", "-b:a", "96k", "-movflags", "+faststart"],
    "avi":  ["-threads", "2", "-c:v", "libxvid", "-q:v", "5", "-c:a", "libmp3lame", "-b:a", "96k"],
    "mov":  ["-threads", "2", "-c:v", "libx264", "-preset", "ultrafast", "-tune", "fastdecode", "-crf", "23", "-c:a", "aac", "-b:a", "96k", "-movflags", "+faststart"],
    "webm": ["-threads", "2", "-c:v", "libvpx", "-crf", "10", "-b:v", "1M", "-deadline", "realtime", "-cpu-used", "8", "-c:a", "libopus", "-b:a", "96k"],
    "mkv":  ["-threads", "2", "-c:v", "libx264", "-preset", "ultrafast", "-tune", "fastdecode", "-crf", "23", "-c:a", "aac", "-b:a", "96k"],
}

def _convert_video(file, task_id, fmt):
    in_path = _save_temp(file, os.path.splitext(file.filename)[1] or ".mp4")
    out_path = in_path.rsplit(".", 1)[0] + "_out." + fmt
    in_ext = os.path.splitext(file.filename)[1]
    fmt_args = FFMPEG_VIDEO_FORMATS[fmt]
    try:
        if task_id: _set_progress(task_id, 3, '分析视频...')
        if _is_same_codec_family(in_ext, fmt):
            # Stream copy - no re-encode needed!
            if task_id: _set_progress(task_id, 10, '快速封装中(无需转码)...')
            _run_ffmpeg(["-i", in_path, "-c", "copy", "-movflags", "+faststart", out_path])
        else:
            dur = _get_duration(in_path)
            if task_id: _set_progress(task_id, 5, '开始转码...')
            _run_ffmpeg(["-i", in_path] + fmt_args + [out_path], task_id=task_id, duration=dur)
        mimes = {"mp4": "video/mp4", "avi": "video/x-msvideo", "mov": "video/quicktime", "webm": "video/webm", "mkv": "video/x-matroska"}
        return send_file(out_path, as_attachment=True,
                        download_name=_download_name(file.filename, "." + fmt),
                        mimetype=mimes[fmt])
    finally:
        _safe_unlink(in_path)
        _safe_unlink(out_path)

# ── Spreadsheet Converters ──

def _csv2xlsx(file, task_id):
    import pandas as pd
    df = pd.read_csv(file.stream)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Sheet1")
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                    download_name=_download_name(file.filename, ".xlsx"),
                    mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

def _xlsx2csv(file, task_id):
    import pandas as pd
    df = pd.read_excel(file.stream, engine="openpyxl")
    buf = io.StringIO()
    df.to_csv(buf, index=False, encoding="utf-8-sig")
    buf.seek(0)
    return send_file(io.BytesIO(buf.getvalue().encode("utf-8-sig")), as_attachment=True,
                    download_name=_download_name(file.filename, ".csv"),
                    mimetype="text/csv")


@app.route("/download/<token>")
def download_file(token):
    token = os.path.basename(token)
    dest = os.path.join(DOWNLOAD_DIR, token)
    metafile = dest + ".json"
    if not os.path.exists(dest) or not os.path.exists(metafile):
        return jsonify({"error": "not found"}), 404
    with open(metafile) as f:
        meta = json.load(f)
    from urllib.parse import quote
    resp = __import__("flask").make_response("")
    resp.headers["Content-Type"] = meta.get("mime", "application/octet-stream")
    resp.headers["Content-Disposition"] = "attachment; filename=" + meta["name"]
    resp.headers["X-Accel-Redirect"] = "/xfile/" + token
    return resp

@app.after_request
def add_security_headers(response):
    response.headers["Server"] = ""
    return response

if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=5000)






